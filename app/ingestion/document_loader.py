"""
ingestion/document_loader.py

Loads PDF and DOCX files and returns a list of LangChain Document objects.
Supports both text extraction and basic table-aware extraction via pdfplumber.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ─── PDF Loading ──────────────────────────────────────────────────────────────

def _load_pdf(file_path: Path) -> List[Document]:
    """Extract text from a PDF file page by page using PyMuPDF (fitz)."""
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError("Install pymupdf: pip install pymupdf")

    documents: List[Document] = []
    doc = fitz.open(str(file_path))

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if not text:
            logger.debug("Page %d is empty or image-only, skipping.", page_num)
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "page": page_num,
                    "file_type": "pdf",
                },
            )
        )

    doc.close()
    logger.info("Loaded %d pages from PDF: %s", len(documents), file_path.name)
    return documents


def _load_pdf_with_tables(file_path: Path) -> List[Document]:
    """
    Fallback PDF loader using pdfplumber — better for table-heavy documents.
    Called automatically when PyMuPDF yields very little text.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip install pdfplumber")

    documents: List[Document] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text_parts: list[str] = []

            # Plain text
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(text.strip())

            # Tables → plain text rows
            for table in page.extract_tables():
                rows = [
                    " | ".join(str(cell) if cell else "" for cell in row)
                    for row in table
                    if row
                ]
                if rows:
                    text_parts.append("\n".join(rows))

            combined = "\n\n".join(text_parts).strip()
            if not combined:
                continue

            documents.append(
                Document(
                    page_content=combined,
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "page": page_num,
                        "file_type": "pdf",
                        "loader": "pdfplumber",
                    },
                )
            )

    logger.info(
        "Loaded %d pages from PDF (pdfplumber): %s", len(documents), file_path.name
    )
    return documents


# ─── DOCX Loading ─────────────────────────────────────────────────────────────

def _load_docx(file_path: Path) -> List[Document]:
    """Extract text from a DOCX file, grouping by paragraph blocks."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    docx = DocxDocument(str(file_path))
    full_text_parts: list[str] = []

    # Paragraphs
    for para in docx.paragraphs:
        text = para.text.strip()
        if text:
            full_text_parts.append(text)

    # Tables
    for table in docx.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text_parts.append(row_text)

    full_text = "\n\n".join(full_text_parts)

    if not full_text.strip():
        logger.warning("No text extracted from DOCX: %s", file_path.name)
        return []

    documents = [
        Document(
            page_content=full_text,
            metadata={
                "source": str(file_path),
                "filename": file_path.name,
                "page": 1,
                "file_type": "docx",
            },
        )
    ]
    logger.info("Loaded DOCX: %s (%d chars)", file_path.name, len(full_text))
    return documents


# ─── Public Interface ─────────────────────────────────────────────────────────

def load_document(file_path: str | Path) -> List[Document]:
    """
    Load a PDF or DOCX file and return a list of LangChain Documents.

    Args:
        file_path: Path to the uploaded file.

    Returns:
        List of Document objects with page_content and metadata.

    Raises:
        ValueError: If the file type is unsupported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        docs = _load_pdf(path)
        # If very little text was extracted, retry with pdfplumber
        total_chars = sum(len(d.page_content) for d in docs)
        if total_chars < 200:
            logger.warning(
                "PyMuPDF extracted very little text (%d chars). "
                "Retrying with pdfplumber.",
                total_chars,
            )
            docs = _load_pdf_with_tables(path)
        return docs

    elif suffix in (".docx", ".doc"):
        return _load_docx(path)

    else:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. Only PDF and DOCX are supported."
        )
